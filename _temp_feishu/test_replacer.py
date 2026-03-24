#!/usr/bin/env python3
"""测试图片提取和 URL 替换逻辑"""

import re
import zipfile
import tempfile
import os

# 测试数据
test_md_content = """# Test Document

![image](https://internal-api-drive-stream.feishu.com/space/api/box/stream/download/v2/cover/TOKEN1)

Some text here.

![image](https://internal-api-drive-stream.feishu.com/space/api/box/stream/download/v2/cover/TOKEN2)

More text.

![image](https://internal-api-drive-stream.feishu.com/space/api/box/stream/download/v2/cover/TOKEN3)
"""

# 测试 URL 提取
pattern = r'!\[image\]\(https://internal-api-drive-stream\.feishu\.com/space/api/box/stream/download/v2/cover/[^)]+\)'
urls = re.findall(pattern, test_md_content)
print(f"Found {len(urls)} URLs:")
for u in urls:
    print(f"  {u}")

# 测试 URL 替换
qiniu_urls = [
    "https://synbiopath.online/feishu_docx/image_1.png",
    "https://synbiopath.online/feishu_docx/image_2.png",
    "https://synbiopath.online/feishu_docx/image_3.png"
]

def replace_func(match):
    if qiniu_urls:
        url = qiniu_urls.pop(0)
        return f'![]({url})'
    return match.group(0)

result = re.sub(pattern, replace_func, test_md_content)
print("\n--- Replaced Markdown ---")
print(result)

# 测试从 DOCX 提取图片
print("\n--- Testing DOCX extraction ---")

# 创建一个简单的带图片的 DOCX 测试
from docx import Document
from docx.shared import Inches

with tempfile.TemporaryDirectory() as tmpdir:
    # 创建测试 DOCX
    test_docx = os.path.join(tmpdir, "test.docx")
    doc = Document()
    doc.add_heading("Test with Images", 0)
    doc.add_paragraph("This is a test paragraph.")
    doc.save(test_docx)

    # 提取图片
    with zipfile.ZipFile(test_docx, 'r') as zf:
        media_files = sorted([n for n in zf.namelist() if 'word/media/' in n])
        print(f"Media files in DOCX: {media_files}")

print("\nTest completed successfully!")
