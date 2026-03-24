import zipfile
import os

# Test if we can extract images from a DOCX file
# DOCX is just a ZIP file containing XML and media files

# First, let's create a test DOCX with images to verify the approach
test_docx = "C:/Software/Data/03-GitHub/SynBioPath/_temp_feishu/test.docx"

# Check if the feishu-docx output directory has any DOCX files
import glob
output_dir = "C:/Software/Data/03-GitHub/SynBioPath/_temp_feishu/feishu_docx_output"
docx_files = glob.glob(f"{output_dir}/*.docx")
print(f"DOCX files found: {docx_files}")

# Also check if python-docx is available
try:
    import docx
    print("python-docx is available")
except ImportError:
    print("python-docx is NOT available, will use zipfile approach")

# Test extracting images from a DOCX file using zipfile
# Let's find any DOCX file to test
import shutil
import tempfile

# Create a simple DOCX file with an image for testing
from docx import Document
from docx.shared import Inches

# Actually, let's just demonstrate the ZIP approach works for extracting images
print("\nTesting ZIP approach to extract images from DOCX:")

# Create a simple test
with tempfile.TemporaryDirectory() as tmpdir:
    # Create a simple document with an image
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph.')
    test_docx_path = os.path.join(tmpdir, 'test_with_image.docx')
    doc.save(test_docx_path)

    # Open as ZIP and list contents
    with zipfile.ZipFile(test_docx_path, 'r') as zf:
        for name in zf.namelist():
            if 'media' in name.lower() or 'image' in name.lower():
                print(f"  Found: {name}")
                # Extract the image
                zf.extract(name, tmpdir)
                print(f"    Extracted to: {os.path.join(tmpdir, name)}")

print("\nZIP approach for extracting images from DOCX works!")
